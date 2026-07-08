#!/usr/bin/env python3
from __future__ import annotations

import argparse
import fnmatch
import json
import os
import re
from datetime import datetime
from zoneinfo import ZoneInfo
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SKIP_DIRS = {
    ".git", ".venv", "venv", "node_modules", "target", "build",
    "dist", "coverage", "cdk.out", "__pycache__", "archive"
}

AWS_RESOURCE_PATTERNS = {
    "ECS Cluster": [("ecs.Cluster", r"\becs\.Cluster\s*\(")],
    "ECS Fargate Service": [("ecs.FargateService", r"\becs\.FargateService\s*\(")],
    "ECS Task Definition": [("ecs.FargateTaskDefinition", r"\becs\.FargateTaskDefinition\s*\(")],
    "ECR Repository": [("ecr.Repository", r"\becr\.Repository\b")],
    "Application Load Balancer": [("elbv2.ApplicationLoadBalancer", r"\belbv2\.ApplicationLoadBalancer\s*\(")],
    "Network Load Balancer": [("elbv2.NetworkLoadBalancer", r"\belbv2\.NetworkLoadBalancer\s*\(")],
    "Target Group": [("elbv2.TargetGroup/add_targets", r"\b(add_targets|ApplicationTargetGroup|NetworkTargetGroup)\s*\(")],
    "Listener Rule": [("elbv2.ApplicationListenerRule", r"\belbv2\.ApplicationListenerRule\s*\(")],
    "VPC": [("ec2.Vpc", r"\bec2\.Vpc(\.from_lookup)?\s*\(")],
    "Security Group": [("ec2.SecurityGroup", r"\bec2\.SecurityGroup\s*\(")],
    "Secrets Manager Secret": [("secretsmanager.Secret", r"\bsecretsmanager\.Secret(\.from_secret_name_v2)?\s*\(")],
    "S3 Bucket": [("s3.Bucket", r"\bs3\.Bucket(\.from_bucket_name)?\s*\(")],
    "CloudFront Distribution": [("cloudfront.Distribution", r"\bcloudfront\.Distribution\s*\(")],
    "CloudFront Key": [("cloudfront.PublicKey/KeyGroup", r"\bcloudfront\.(PublicKey|KeyGroup)\s*\(")],
    "OpenSearch Domain": [("opensearch.Domain", r"\bopensearch\.Domain\s*\(")],
    "KMS Key": [("kms.Key", r"\bkms\.Key\s*\(")],
    "EFS File System": [("efs.FileSystem", r"\befs\.FileSystem\s*\(")],
    "EFS Access Point": [("efs.AccessPoint/add_access_point", r"\badd_access_point\s*\(")],
    "IAM Policy/Role": [("iam.PolicyStatement/Role", r"\biam\.(PolicyStatement|Role)\s*\(")],
    "RDS": [("rds.DatabaseCluster/Instance", r"\brds\.(DatabaseCluster|DatabaseInstance)\s*\(")],
    "Lambda Function": [("lambda.Function", r"\blambda_?\.Function\s*\(")],
    "SQS Queue": [("sqs.Queue", r"\bsqs\.Queue\s*\(")],
    "SNS Topic": [("sns.Topic", r"\bsns\.Topic\s*\(")],
    "DynamoDB Table": [("dynamodb.Table", r"\bdynamodb\.Table\s*\(")],
    "API Gateway": [("apigateway/apigatewayv2", r"\b(apigateway|apigatewayv2)\.")],
    "Route53": [("route53", r"\broute53\.")],
    "ACM Certificate": [("acm.Certificate", r"\bacm\.Certificate\b")],
    "CloudWatch": [("cloudwatch", r"\bcloudwatch\.")],
    "EventBridge Rule": [("events.Rule", r"\bevents\.Rule\s*\(")]
}


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate SOP DOCX from project metadata and repo/CDK scan.")
    parser.add_argument("--config", required=True)
    parser.add_argument("--repo", default=".")
    parser.add_argument("--out", required=True)
    parser.add_argument("--cdk-dir", default="awscdk")
    parser.add_argument("--print-discovery", action="store_true")
    args = parser.parse_args()

    repo = Path(args.repo).resolve()
    config = json.loads(Path(args.config).read_text(encoding="utf-8-sig"))
    discovered = scan_repo(repo, args.cdk_dir, config)

    if args.print_discovery:
        print(json.dumps(discovered, indent=2))
        return 0

    out = Path(args.out).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(config, discovered, out)
    print(f"Wrote SOP: {out}")
    return 0


def scan_repo(repo: Path, cdk_dir: str, config: dict[str, Any]) -> dict[str, Any]:
    dockerfiles = find_files(repo, ["**/Dockerfile", "**/*.dockerfile"])
    workflows = find_files(repo, [".github/workflows/*.yml", ".github/workflows/*.yaml"])
    runtime_files = find_files(repo, ["**/package.json", "**/pom.xml", "**/requirements*.txt", "**/pyproject.toml", "**/go.mod"])
    cdk_root = repo / cdk_dir
    cdk_files = find_files(cdk_root, ["**/*.py", "**/*.ts"]) if cdk_root.exists() else []
    services = merge_by_name(config.get("services", []), discover_services(repo, dockerfiles, cdk_files))
    aws_resources = discover_aws_resources(repo, cdk_files)

    return {
        "repo": str(repo),
        "services": services,
        "dockerfiles": [rel(repo, p) for p in dockerfiles],
        "runtime_files": [summarize_runtime(repo, p) for p in runtime_files],
        "workflows": [summarize_workflow(repo, p) for p in workflows],
        "aws_resources": aws_resources,
        "aws_resource_names": sorted({resource["resource"] for resource in aws_resources}),
        "cdk_files": [rel(repo, p) for p in cdk_files],
        "environments": merge_environments(config.get("environments", []), discover_envs(workflows), config)
    }


def find_files(root: Path, patterns: list[str]) -> list[Path]:
    if not root.exists():
        return []

    found: list[Path] = []
    for current, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS]
        current_path = Path(current)

        for filename in filenames:
            path = current_path / filename
            rel_path = path.relative_to(root).as_posix()
            if any(pattern_matches(rel_path, filename, pattern) for pattern in patterns):
                found.append(path)

    return sorted(found)


def pattern_matches(rel_path: str, filename: str, pattern: str) -> bool:
    if fnmatch.fnmatch(rel_path, pattern) or fnmatch.fnmatch(filename, pattern):
        return True
    if pattern.startswith("**/") and fnmatch.fnmatch(rel_path, pattern[3:]):
        return True
    return False


def discover_services(repo: Path, dockerfiles: list[Path], cdk_files: list[Path]) -> list[dict[str, Any]]:
    services: dict[str, dict[str, Any]] = {}

    for dockerfile in dockerfiles:
        name = dockerfile.parent.name
        item = services.setdefault(name, {"name": name})
        item["dockerfile"] = rel(repo, dockerfile)
        item.setdefault("runtime", infer_runtime(dockerfile.parent))
        item.setdefault("artifact", "Docker image")

    for path in cdk_files:
        rel_path = rel(repo, path)
        if "/services/" in f"/{rel_path}/":
            name = path.stem
            if name not in {"__init__", "base", "base_service"} and not name.startswith("_"):
                item = services.setdefault(name, {"name": name})
                item["cdk_service_file"] = rel_path
                item.setdefault("artifact", "Docker image")

    return list(services.values())


def merge_by_name(config_items: list[dict[str, Any]], discovered_items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged = {item["name"]: dict(item) for item in discovered_items if item.get("name")}

    for item in config_items:
        name = item.get("name")
        if name:
            merged.setdefault(name, {"name": name}).update({
                key: value for key, value in item.items()
                if value not in ("", None)
            })

    return sorted(merged.values(), key=lambda item: item["name"])


def infer_runtime(folder: Path) -> str:
    if (folder / "package.json").exists():
        return "Node.js"
    if (folder / "pom.xml").exists():
        return "Java/Maven"
    if (folder / "requirements.txt").exists() or (folder / "pyproject.toml").exists():
        return "Python"
    if (folder / "go.mod").exists():
        return "Go"
    return ""


def discover_aws_resources(repo: Path, cdk_files: list[Path]) -> list[dict[str, str]]:
    found: dict[tuple[str, str, str, str], dict[str, str]] = {}

    for path in cdk_files:
        text = read(path)
        for resource, patterns in AWS_RESOURCE_PATTERNS.items():
            for construct, pattern in patterns:
                for match in re.finditer(pattern, text):
                    construct_id = guess_construct_id(text, match.start())
                    key = (resource, construct, construct_id, rel(repo, path))
                    found[key] = {
                        "resource": resource,
                        "construct": construct,
                        "construct_id": construct_id or "",
                        "source": rel(repo, path)
                    }

    return sorted(found.values(), key=lambda item: (item["resource"], item["source"], item["construct_id"]))


def guess_construct_id(text: str, start: int) -> str:
    snippet = text[start:start + 600]
    patterns = [
        r"\(\s*self\s*,\s*[\"']([^\"']+)[\"']",
        r"\(\s*stack\s*,\s*[\"']([^\"']+)[\"']",
        r"\(\s*scope\s*,\s*[\"']([^\"']+)[\"']",
        r"\(\s*[\"']([^\"']+)[\"']",
        r"id\s*=\s*[\"']([^\"']+)[\"']"
    ]

    for pattern in patterns:
        match = re.search(pattern, snippet, re.DOTALL)
        if match:
            return match.group(1)

    return ""


def summarize_runtime(repo: Path, path: Path) -> dict[str, str]:
    text = read(path)

    if path.name == "package.json":
        try:
            package = json.loads(text)
            return {
                "file": rel(repo, path),
                "type": "Node.js",
                "name": package.get("name", ""),
                "details": ", ".join((package.get("scripts") or {}).keys())
            }
        except json.JSONDecodeError:
            return {"file": rel(repo, path), "type": "Node.js", "name": "", "details": "Invalid package.json"}

    if path.name == "pom.xml":
        return {
            "file": rel(repo, path),
            "type": "Java/Maven",
            "name": first(text, r"<artifactId>([^<]+)</artifactId>"),
            "details": first(text, r"<java.version>([^<]+)</java.version>")
        }

    if path.name.startswith("requirements"):
        deps = [line for line in text.splitlines() if line.strip() and not line.strip().startswith("#")]
        return {
            "file": rel(repo, path),
            "type": "Python requirements",
            "name": "",
            "details": f"{len(deps)} dependencies"
        }

    return {"file": rel(repo, path), "type": path.name, "name": "", "details": ""}


def summarize_workflow(repo: Path, path: Path) -> dict[str, Any]:
    text = read(path)
    name = first(text, r"(?m)^name:\s*(.+)$") or path.stem
    lower = f"{path.name} {name}".lower()

    if "deploy" in lower:
        category = "deploy"
    elif "release" in lower:
        category = "release"
    elif "build" in lower or "image" in lower:
        category = "build"
    else:
        category = "other"

    return {
        "file": rel(repo, path),
        "name": name,
        "category": category,
        "triggers": sorted(set(re.findall(r"(?m)^\s{0,4}(workflow_dispatch|push|pull_request|schedule):", text))),
        "jobs": sorted(set(re.findall(r"(?m)^  ([A-Za-z0-9_-]+):\s*$", text)))
    }


def discover_envs(workflows: list[Path]) -> list[str]:
    envs = {"dev", "dev2", "qa", "qa2", "stage", "prod"}
    found = set()

    for workflow in workflows:
        text = read(workflow)
        found.update(env for env in envs if re.search(rf"\b{env}\b", text))

    return sorted(found, key=env_sort_key)


def merge_environments(config_envs: list[dict[str, Any]], discovered: list[str], config: dict[str, Any]) -> list[dict[str, str]]:
    merged = {
        name: {"name": name, "purpose": "", "url": "", "notes": "Discovered from workflows"}
        for name in discovered
    }

    for env in config_envs:
        if env.get("name"):
            merged.setdefault(env["name"], {
                "name": env["name"],
                "purpose": "",
                "url": "",
                "notes": ""
            }).update({key: str(value) for key, value in env.items()})

    for env in merged.values():
        if not env.get("url") and config.get("url_template"):
            if env["name"] == "prod" and config.get("prod_url"):
                env["url"] = config["prod_url"]
            else:
                env["url"] = config["url_template"].replace("{env}", env["name"])

    return sorted(merged.values(), key=lambda item: env_sort_key(item["name"]))


def build_docx(config: dict[str, Any], discovered: dict[str, Any], out: Path) -> None:
    doc = Document()
    setup_doc(doc)
    project = config.get("project", {})

    title(doc, f"{project.get('name', 'Project')} SOP")
    paragraph(doc, f"Generated {datetime.now(ZoneInfo('America/New_York')).strftime('%Y-%m-%d %H:%M %Z')}", italic=True, center=True)

    heading(doc, "1. Project Overview")
    labeled_bullets(doc, [
        ("Project name", project.get("name", "TBD")),
        ("Description", project.get("description", "TBD")),
        ("Owner/team", project.get("owner", "TBD")),
        ("Primary contact", project.get("primary_contact", "TBD")),
        ("Support channel", project.get("support_channel", "TBD"))
    ])

    heading(doc, "2. Services")
    table(doc, ["Service", "Description", "Runtime", "Dockerfile", "Artifact"], [
        [
            service.get("name", ""),
            service.get("description", "TBD"),
            service.get("runtime", "TBD"),
            service.get("dockerfile", "TBD"),
            service.get("artifact", "TBD")
        ]
        for service in discovered["services"]
    ])

    heading(doc, "3. AWS Resources")
    table(doc, ["AWS Resource"], [[name] for name in discovered.get("aws_resource_names", [])])

    heading(doc, "4. Dockerfiles")
    simple_bullets(doc, discovered["dockerfiles"] or ["TBD"])

    heading(doc, "5. Environments And URLs")
    table(doc, ["Environment", "Purpose", "URL", "Notes"], [
        [env["name"], env.get("purpose", ""), env.get("url", ""), env.get("notes", "")]
        for env in discovered["environments"]
    ])

    heading(doc, "6. Build And Deploy Workflows")
    table(doc, ["Category", "Workflow", "File"], [
        [
            workflow["category"],
            workflow["name"],
            workflow["file"]
        ]
        for workflow in discovered["workflows"]
    ])

    for number, key, label in [
        (7, "build", "Build Procedure"),
        (8, "deploy", "Deployment Procedure"),
        (9, "verify", "Verification")
    ]:
        steps = config.get("runbooks", {}).get(key)
        if steps:
            heading(doc, f"{number}. {label}")
            numbered(doc, steps)

    doc.save(out)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.75)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)


def title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)


def heading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def subheading(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def paragraph(doc: Document, text: str, italic: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic


def labeled_bullets(doc: Document, items: list[tuple[str, Any]]) -> None:
    for label, value in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value or "TBD"))


def simple_bullets(doc: Document, items: list[Any]) -> None:
    for item in items or ["TBD"]:
        doc.add_paragraph(str(item), style="List Bullet")


def numbered(doc: Document, items: list[Any]) -> None:
    for item in items or ["TBD"]:
        doc.add_paragraph(str(item), style="List Number")


def table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    rows = rows or [["TBD"] + [""] * (len(headers) - 1)]
    doc_table = doc.add_table(rows=1, cols=len(headers))
    doc_table.style = "Table Grid"

    for index, header in enumerate(headers):
        cell = doc_table.rows[0].cells[index]
        cell.text = header
        for paragraph_obj in cell.paragraphs:
            for run in paragraph_obj.runs:
                run.bold = True

    for row in rows:
        cells = doc_table.add_row().cells
        for index, value in enumerate(row[:len(headers)]):
            cells[index].text = str(value or "")

    doc.add_paragraph()


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def rel(root: Path, path: Path) -> str:
    try:
        return str(path.relative_to(root)).replace(os.sep, "/")
    except ValueError:
        return str(path)


def first(text: str, pattern: str) -> str:
    match = re.search(pattern, text)
    return match.group(1).strip().strip("'\"") if match else ""


def env_sort_key(name: str) -> tuple[int, str]:
    order = ["dev", "dev2", "qa", "qa2", "stage", "prod"]
    return (order.index(name) if name in order else 99, name)


if __name__ == "__main__":
    raise SystemExit(main())
